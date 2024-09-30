from docx import Document
from docx.shared import Pt,Cm
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
from model import KhachSiQuanModel,KhachQNCNModel,KhachVienChucModel,KhachTheoDoanModel
from django.db.models import Q
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from django.utils import timezone
from datetime import timedelta
# from spire.doc import *
# from spire.doc.common import *



def create_word_document():
    # Tạo tài liệu Word mới
    cochu =14
    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Cm(29.7)  # Chiều rộng của khổ A4 nằm ngang
        section.page_height = Cm(21.0)  # Chiều cao của khổ A4 nằm ngang

        # Đặt lề của tài liệu
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.top_margin = Cm(1.5)

    # Thêm phần "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM"
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    
    # Đặt font chữ và cỡ chữ
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Căn giữa đoạn văn
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Thiết lập giãn dòng 1.3
    paragraph.paragraph_format.line_spacing = Pt(14)  # Đặt giãn dòng bằng cỡ chữ, có thể điều chỉnh

    # Xóa khoảng cách giữa các đoạn
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    # Thêm dòng "Độc lập - Tự do - Hạnh phúc" và gạch dưới
    paragraph = doc.add_paragraph()
    run = paragraph.add_run("Độc lập - Tự do - Hạnh phúc")
    
    # Đặt font chữ và cỡ chữ
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Gạch dưới
    run.underline = True
    
    # Căn giữa đoạn văn
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Thiết lập giãn dòng 1.3
    paragraph.paragraph_format.line_spacing = 1.3

    now = datetime.now()
    day = now.day
    month = now.strftime('%m')  # Thay đổi định dạng tháng nếu cần
    year = now.year

    # Thêm dòng chữ "Hà Nội, ngày [Ngày] tháng [Tháng] năm [Năm]" căn phải
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"Hà Nội, ngày {day} tháng {month} năm {year}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    # Căn phải đoạn văn
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Thêm nội dung vào tài liệu
    paragraph = doc.add_paragraph()
    run_number = paragraph.add_run("1. ")
    run_number.font.name = 'Times New Roman'
    run_number.font.size = Pt(14)
    run_number.bold = True  # In đậm

    # Thêm nội dung "Danh sách vẫn đang tiếp khách" và in đậm
    run_text = paragraph.add_run("Danh sách vẫn đang tiếp khách của sĩ quan")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(14)
    paragraph.paragraph_format.line_spacing = 1.3
    run_text.bold = True  # In đậm``
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    # table.autofit = False
    table.autofit = False 
    # table.allow_autofit = False

    # Điều chỉnh độ rộng các cột theo yêu cầu
    # table.columns[0].width = Cm(4)   # Cột 1: 5 cm
    # table.columns[1].width = Cm(4)   # Cột 2: 5 cm
    # table.columns[2].width = Cm(4)   # Cột 3: 5 cm
    # table.columns[3].width = Cm(3)   # Cột 4: 4 cm
    # table.columns[4].width = Cm(4)   # Cột 5: 4 cm
    # table.columns[5].width = Cm(1.5) # Cột 6: 1.5 cm
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Thêm tiêu đề cột
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên sĩ quan'
    hdr_cells[1].text = 'Đơn vị'
    hdr_cells[2].text = 'Tên khách'
    hdr_cells[3].text = 'Số định danh'
    hdr_cells[4].text = 'Thời gian bắt đầu'
    hdr_cells[5].text = 'Số thẻ'
    hdr_cells[6].text = 'Ghi chú'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    khach_data = get_khach_dang_tiep_don_sq()

    # Duyệt qua danh sách khách để thêm vào bảng
    for khach in khach_data:
        row_cells = table.add_row().cells
        row_cells[0].text = khach['TenSiQuan']
        row_cells[1].text = khach['TenDonViSiQuan']
        row_cells[2].text = khach['TenKhach']
        row_cells[3].text = khach['SoDinhDanh']
        row_cells[4].text = khach['ThoiGianBatDau'].strftime("%H:%M %d/%m/%Y")
        row_cells[5].text = khach['SoTheKhach']
        row_cells[6].text = khach.get('GhiChu', '')

        for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

    for row in table.rows:
            # row.cells[idx].width = width
            row.cells[0].width = Cm(4)   # Cột 1: 5 cm
            row.cells[1].width = Cm(2.5)   # Cột 2: 5 cm
            row.cells[2].width = Cm(4)   # Cột 3: 5 cm
            row.cells[3].width = Cm(3)   # Cột 4: 4 cm
            row.cells[4].width = Cm(4)   # Cột 5: 4 cm
            row.cells[5].width = Cm(1.8) # Cột 6: 1.5 cm
            
            # row.cells[6].width = Cm(3)
    # Lưu tài liệu
    # doc.save('DanhSachKhachDangTiep.docx')

# Ví dụ dữ liệu khách đang tiếp khách
    doc.add_paragraph()  # Thêm một đoạn văn trống
    # paragraph.paragraph_format.space_after = Pt(0.3)
    paragraph = doc.add_paragraph()
    run_number = paragraph.add_run("2. ")
    run_number.font.name = 'Times New Roman'
    run_number.font.size = Pt(14)
    run_number.bold = True  # In đậm

    # Thêm nội dung "Danh sách vẫn đang tiếp khách" và in đậm
    run_text = paragraph.add_run("Danh sách vẫn đang tiếp khách của QNCN")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(14)
    paragraph.paragraph_format.line_spacing = 1.3
    run_text.bold = True  # In đậm``
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    # table.autofit = False
    table.autofit = False 
    # table.allow_autofit = False

    # Điều chỉnh độ rộng các cột theo yêu cầu
    # table.columns[0].width = Cm(4)   # Cột 1: 5 cm
    # table.columns[1].width = Cm(4)   # Cột 2: 5 cm
    # table.columns[2].width = Cm(4)   # Cột 3: 5 cm
    # table.columns[3].width = Cm(3)   # Cột 4: 4 cm
    # table.columns[4].width = Cm(4)   # Cột 5: 4 cm
    # table.columns[5].width = Cm(1.5) # Cột 6: 1.5 cm
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Thêm tiêu đề cột
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên quân nhân'
    hdr_cells[1].text = 'Đơn vị'
    hdr_cells[2].text = 'Tên khách'
    hdr_cells[3].text = 'Số định danh'
    hdr_cells[4].text = 'Thời gian bắt đầu'
    hdr_cells[5].text = 'Số thẻ'
    hdr_cells[6].text = 'Ghi chú'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    khach_data = get_khach_dang_tiep_don_qncn()

    # Duyệt qua danh sách khách để thêm vào bảng
    for khach in khach_data:
        row_cells = table.add_row().cells
        row_cells[0].text = khach['TenSiQuan']
        row_cells[1].text = khach['TenDonViSiQuan']
        row_cells[2].text = khach['TenKhach']
        row_cells[3].text = khach['SoDinhDanh']
        row_cells[4].text = khach['ThoiGianBatDau'].strftime("%H:%M %d/%m/%Y")
        row_cells[5].text = khach['SoTheKhach']
        row_cells[6].text = khach.get('GhiChu', '')

        for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

    for row in table.rows:
            # row.cells[idx].width = width
            row.cells[0].width = Cm(4)   # Cột 1: 5 cm
            row.cells[1].width = Cm(2.5)   # Cột 2: 5 cm
            row.cells[2].width = Cm(4)   # Cột 3: 5 cm
            row.cells[3].width = Cm(3)   # Cột 4: 4 cm
            row.cells[4].width = Cm(4)   # Cột 5: 4 cm
            row.cells[5].width = Cm(1.8) # Cột 6: 1.5 cm
    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    run_number = paragraph.add_run("3. ")
    run_number.font.name = 'Times New Roman'
    run_number.font.size = Pt(14)
    run_number.bold = True  # In đậm

    # Thêm nội dung "Danh sách vẫn đang tiếp khách" và in đậm
    run_text = paragraph.add_run("Danh sách vẫn đang tiếp khách của viên chức quốc phòng")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(14)
    paragraph.paragraph_format.line_spacing = 1.3
    run_text.bold = True  # In đậm``
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    # table.autofit = False
    table.autofit = False 
    # table.allow_autofit = False

    # Điều chỉnh độ rộng các cột theo yêu cầu
    # table.columns[0].width = Cm(4)   # Cột 1: 5 cm
    # table.columns[1].width = Cm(4)   # Cột 2: 5 cm
    # table.columns[2].width = Cm(4)   # Cột 3: 5 cm
    # table.columns[3].width = Cm(3)   # Cột 4: 4 cm
    # table.columns[4].width = Cm(4)   # Cột 5: 4 cm
    # table.columns[5].width = Cm(1.5) # Cột 6: 1.5 cm
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Thêm tiêu đề cột
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên viên chức'
    hdr_cells[1].text = 'Đơn vị'
    hdr_cells[2].text = 'Tên khách'
    hdr_cells[3].text = 'Số định danh'
    hdr_cells[4].text = 'Thời gian bắt đầu'
    hdr_cells[5].text = 'Số thẻ'
    hdr_cells[6].text = 'Ghi chú'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    khach_data = get_khach_dang_tiep_don_vc()

    # Duyệt qua danh sách khách để thêm vào bảng
    for khach in khach_data:
        row_cells = table.add_row().cells
        row_cells[0].text = khach['TenSiQuan']
        row_cells[1].text = khach['TenDonViSiQuan']
        row_cells[2].text = khach['TenKhach']
        row_cells[3].text = khach['SoDinhDanh']
        row_cells[4].text = khach['ThoiGianBatDau'].strftime("%H:%M %d/%m/%Y")
        row_cells[5].text = khach['SoTheKhach']
        row_cells[6].text = khach.get('GhiChu', '')

        for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

    for row in table.rows:
            # row.cells[idx].width = width
            row.cells[0].width = Cm(4)   # Cột 1: 5 cm
            row.cells[1].width = Cm(2.5)   # Cột 2: 5 cm
            row.cells[2].width = Cm(4)   # Cột 3: 5 cm
            row.cells[3].width = Cm(3)   # Cột 4: 4 cm
            row.cells[4].width = Cm(4)   # Cột 5: 4 cm
            row.cells[5].width = Cm(1.8) # Cột 6: 1.5 cm
    doc.add_paragraph()
    paragraph = doc.add_paragraph()
    run_number = paragraph.add_run("4. ")
    run_number.font.name = 'Times New Roman'
    run_number.font.size = Pt(14)
    run_number.bold = True  # In đậm

    # Thêm nội dung "Danh sách vẫn đang tiếp khách" và in đậm
    run_text = paragraph.add_run("Danh sách trả khách trong ngày của sĩ quan")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(14)
    run_text.bold = True  # In đậm``
    paragraph.paragraph_format.line_spacing = 1.3
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    # table.autofit = False
    table.autofit = False 
    # table.allow_autofit = False

    # Điều chỉnh độ rộng các cột theo yêu cầu
    # table.columns[0].width = Cm(4)   # Cột 1: 5 cm
    # table.columns[1].width = Cm(4)   # Cột 2: 5 cm
    # table.columns[2].width = Cm(4)   # Cột 3: 5 cm
    # table.columns[3].width = Cm(3)   # Cột 4: 4 cm
    # table.columns[4].width = Cm(5)   # Cột 5: 4 cm
    # table.columns[5].width = Cm(1.5) # Cột 6: 1.5 cm
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Thêm tiêu đề cột
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên sĩ quan'
    hdr_cells[1].text = 'Đơn vị'
    hdr_cells[2].text = 'Tên khách'
    hdr_cells[3].text = 'Số định danh'
    hdr_cells[4].text = 'Thời gian bắt đầu'
    hdr_cells[5].text = 'Thời gian kết thúc'
    hdr_cells[6].text = 'Số thẻ'
    hdr_cells[7].text = 'Ghi chú'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    khach_data = get_khach_sq()

    # Duyệt qua danh sách khách để thêm vào bảng
    for khach in khach_data:
        row_cells = table.add_row().cells
        row_cells[0].text = khach['TenSiQuan']
        row_cells[1].text = khach['TenDonViSiQuan']
        row_cells[2].text = khach['TenKhach']
        row_cells[3].text = khach['SoDinhDanh']
        row_cells[4].text = khach['ThoiGianBatDau'].strftime("%H:%M %d/%m/%Y")
        row_cells[5].text = khach['ThoiGianKetThuc'].strftime("%H:%M %d/%m/%Y")
        row_cells[6].text = khach['SoTheKhach']
        row_cells[7].text = khach.get('GhiChu', '')

        for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

    for row in table.rows:
            # row.cells[idx].width = width
            row.cells[0].width = Cm(4)   # Cột 1: 5 cm
            row.cells[1].width = Cm(1.9)   # Cột 2: 5 cm
            row.cells[2].width = Cm(4)   # Cột 3: 5 cm
            row.cells[3].width = Cm(3.5)   # Cột 4: 4 cm
            row.cells[4].width = Cm(4.1)   # Cột 5: 4 cm
            row.cells[5].width = Cm(4.1)   # Cột 5: 4 cm
            row.cells[6].width = Cm(1.3) # Cột 6: 1.5 cm
            row.cells[6].width = Cm(2.2) # Cột 6: 1.5 cm
    #khách qncn
            
            # row.cells[6].width = Cm(3)
    # Lưu tài liệu
    # doc.save('DanhSachKhachDangTiep.docx')

# Ví dụ dữ liệu khách đang tiếp khách
    doc.add_paragraph()  # Thêm một đoạn văn trống
    # paragraph.paragraph_format.space_after = Pt(0.3)

    paragraph = doc.add_paragraph()
    run_number = paragraph.add_run("5. ")
    run_number.font.name = 'Times New Roman'
    run_number.font.size = Pt(14)
    run_number.bold = True  # In đậm

    # Thêm nội dung "Danh sách vẫn đang tiếp khách" và in đậm
    run_text = paragraph.add_run("Danh sách trả khách trong ngày của QNCN")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(14)
    run_text.bold = True  # In đậm``
    paragraph.paragraph_format.line_spacing = 1.3
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    # table.autofit = False
    table.autofit = False 
    # table.allow_autofit = False

    # Điều chỉnh độ rộng các cột theo yêu cầu
    # table.columns[0].width = Cm(4)   # Cột 1: 5 cm
    # table.columns[1].width = Cm(4)   # Cột 2: 5 cm
    # table.columns[2].width = Cm(4)   # Cột 3: 5 cm
    # table.columns[3].width = Cm(3)   # Cột 4: 4 cm
    # table.columns[4].width = Cm(5)   # Cột 5: 4 cm
    # table.columns[5].width = Cm(1.5) # Cột 6: 1.5 cm
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Thêm tiêu đề cột
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên quân nhân'
    hdr_cells[1].text = 'Đơn vị'
    hdr_cells[2].text = 'Tên khách'
    hdr_cells[3].text = 'Số định danh'
    hdr_cells[4].text = 'Thời gian bắt đầu'
    hdr_cells[5].text = 'Thời gian kết thúc'
    hdr_cells[6].text = 'Số thẻ'
    hdr_cells[7].text = 'Ghi chú'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    khach_data = get_khach_qncn()

    # Duyệt qua danh sách khách để thêm vào bảng
    for khach in khach_data:
        row_cells = table.add_row().cells
        row_cells[0].text = khach['TenSiQuan']
        row_cells[1].text = khach['TenDonViSiQuan']
        row_cells[2].text = khach['TenKhach']
        row_cells[3].text = khach['SoDinhDanh']
        row_cells[4].text = khach['ThoiGianBatDau'].strftime("%H:%M %d/%m/%Y")
        row_cells[5].text = khach['ThoiGianKetThuc'].strftime("%H:%M %d/%m/%Y")
        row_cells[6].text = khach['SoTheKhach']
        row_cells[7].text = khach.get('GhiChu', '')

        for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

    for row in table.rows:
            # row.cells[idx].width = width
            row.cells[0].width = Cm(4)   # Cột 1: 5 cm
            row.cells[1].width = Cm(1.9)   # Cột 2: 5 cm
            row.cells[2].width = Cm(4)   # Cột 3: 5 cm
            row.cells[3].width = Cm(3.5)   # Cột 4: 4 cm
            row.cells[4].width = Cm(4.1)   # Cột 5: 4 cm
            row.cells[5].width = Cm(4.1)   # Cột 5: 4 cm
            row.cells[6].width = Cm(1.3) # Cột 6: 1.5 cm
            row.cells[6].width = Cm(2.2) # Cột 6: 1.5 cm 
    #khách vcqp
            
            # row.cells[6].width = Cm(3)
    # Lưu tài liệu
    # doc.save('DanhSachKhachDangTiep.docx')

# Ví dụ dữ liệu khách đang tiếp khách
    doc.add_paragraph()  # Thêm một đoạn văn trống
    # paragraph.paragraph_format.space_after = Pt(0.3)

    paragraph = doc.add_paragraph()
    run_number = paragraph.add_run("6. ")
    run_number.font.name = 'Times New Roman'
    run_number.font.size = Pt(14)
    run_number.bold = True  # In đậm

    # Thêm nội dung "Danh sách vẫn đang tiếp khách" và in đậm
    run_text = paragraph.add_run("Danh sách trả khách trong ngày của viên chức quốc phòng")
    run_text.font.name = 'Times New Roman'
    run_text.font.size = Pt(14)
    run_text.bold = True  # In đậm``
    paragraph.paragraph_format.line_spacing = 1.3
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    # table.autofit = False
    table.autofit = False 
    # table.allow_autofit = False

    # Điều chỉnh độ rộng các cột theo yêu cầu
    # table.columns[0].width = Cm(4)   # Cột 1: 5 cm
    # table.columns[1].width = Cm(4)   # Cột 2: 5 cm
    # table.columns[2].width = Cm(4)   # Cột 3: 5 cm
    # table.columns[3].width = Cm(3)   # Cột 4: 4 cm
    # table.columns[4].width = Cm(5)   # Cột 5: 4 cm
    # table.columns[5].width = Cm(1.5) # Cột 6: 1.5 cm
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Thêm tiêu đề cột
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tên viên chức'
    hdr_cells[1].text = 'Đơn vị'
    hdr_cells[2].text = 'Tên khách'
    hdr_cells[3].text = 'Số định danh'
    hdr_cells[4].text = 'Thời gian bắt đầu'
    hdr_cells[5].text = 'Thời gian kết thúc'
    hdr_cells[6].text = 'Số thẻ'
    hdr_cells[7].text = 'Ghi chú'
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    khach_data = get_khach_vc()

    # Duyệt qua danh sách khách để thêm vào bảng
    for khach in khach_data:
        row_cells = table.add_row().cells
        row_cells[0].text = khach['TenSiQuan']
        row_cells[1].text = khach['TenDonViSiQuan']
        row_cells[2].text = khach['TenKhach']
        row_cells[3].text = khach['SoDinhDanh']
        row_cells[4].text = khach['ThoiGianBatDau'].strftime("%H:%M %d/%m/%Y")
        row_cells[5].text = khach['ThoiGianKetThuc'].strftime("%H:%M %d/%m/%Y")
        row_cells[6].text = khach['SoTheKhach']
        row_cells[7].text = khach.get('GhiChu', '')

        for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(14)

    for row in table.rows:
            # row.cells[idx].width = width
            row.cells[0].width = Cm(4)   # Cột 1: 5 cm
            row.cells[1].width = Cm(1.9)   # Cột 2: 5 cm
            row.cells[2].width = Cm(4)   # Cột 3: 5 cm
            row.cells[3].width = Cm(3.5)   # Cột 4: 4 cm
            row.cells[4].width = Cm(4.1)   # Cột 5: 4 cm
            row.cells[5].width = Cm(4.1)   # Cột 5: 4 cm
            row.cells[6].width = Cm(1.3) # Cột 6: 1.5 cm
            row.cells[6].width = Cm(2.2) # Cột 6: 1.5 cm
            # row.cells[6].width = Cm(3)
    # Lưu tài liệu
    # doc.save('DanhSachKhachDangTiep.docx')



    doc.add_paragraph()
# Ví dụ dữ liệu khách đang tiếp khách
    table = doc.add_table(rows=1, cols=2)

    # Đặt chiều cao cho hàng
    for row  in table.rows:
        row .height = Cm(4)  # C
    # Thêm phần footer
  
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element.tcPr
            # tc.left = None
            # tc.top = None
            # tc.right = None
            # tc.bottom = None
            # cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            # cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

    # Đặt nội dung cho các ô
    cell1 = table.cell(0, 0)
    cell2 = table.cell(0, 1)

    # Định dạng nội dung ô
    for cell, text in zip([cell1, cell2], ['Trực ban giao phiên', 'Trực ban nhận phiên']):
        cell.text = text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 1  # Căn giữa
        
        # Thêm khoảng cách giữa các đoạn
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

        # Đặt khoảng cách sau đoạn văn
    paragraph.paragraph_format.space_after = Pt(4)  # Giãn dòng 3 pt

    # Lưu tài liệu ra file
    file_path = f"BaoCao.docx"
    doc.save(file_path)

    return file_path


def get_khach_dang_tiep_don_sq():
    # Lấy danh sách KhachSiQuanModel có ThoiGianKetThuc là None
    khach_dang_tiep = KhachSiQuanModel.KhachSiQuanModel.objects.filter(ThoiGianKetThuc__isnull=True).order_by("-ThoiGianBatDau")
    
    # Chuẩn bị danh sách kết quả
    result = []

    for khach in khach_dang_tiep:
        # Lấy các thông tin cần thiết
        si_quan = khach.SiQuan
        don_vi = si_quan.DonVi
        khach_hang = khach.Khach
        the_khach = khach.TheKhach

        # Thêm vào danh sách kết quả
        result.append({
            'TenSiQuan': si_quan.HoTen,
            'TenDonViSiQuan': don_vi.MaDonVi,
            'TenKhach': khach_hang.HoTenKhach,
            'SoDinhDanh': khach_hang.SoDinhDanh,
            'ThoiGianBatDau': khach.ThoiGianBatDau,
            'SoTheKhach': the_khach.SoThe,
            'GhiChu': khach.GhiChu
        })

    return result


def get_khach_sq():


# Lấy thời gian 16h30 của ngày hôm trước
    now = timezone.now()
    start_of_yesterday = (now - timedelta(days=1)).replace(hour=16, minute=30, second=0, microsecond=0)
    
    # Lọc các bản ghi có ThoiGianKetThuc là None hoặc ThoiGianBatDau từ 16h30 ngày hôm trước
    khach_dang_tiep = KhachSiQuanModel.KhachSiQuanModel.objects.filter(ThoiGianKetThuc__gte=start_of_yesterday).order_by("-ThoiGianBatDau")
    # khach_dang_tiep = KhachSiQuanModel.KhachSiQuanModel.objects.all().order_by("-ThoiGianBatDau")
    # khach_dang_tiep = KhachSiQuanModel.KhachSiQuanModel.objects.filter(ThoiGianKetThuc__isnull=False).order_by("-ThoiGianBatDau")
    # Chuẩn bị danh sách kết quả
    result = []

    for khach in khach_dang_tiep:
        # Lấy các thông tin cần thiết
        si_quan = khach.SiQuan
        don_vi = si_quan.DonVi
        khach_hang = khach.Khach
        the_khach = khach.TheKhach

        # Thêm vào danh sách kết quả
        result.append({
            'TenSiQuan': si_quan.HoTen,
            'TenDonViSiQuan': don_vi.MaDonVi,
            'TenKhach': khach_hang.HoTenKhach,
            'SoDinhDanh': khach_hang.SoDinhDanh,
            'ThoiGianBatDau': khach.ThoiGianBatDau,
            'ThoiGianKetThuc': khach.ThoiGianKetThuc,
            'SoTheKhach': the_khach.SoThe,
            'GhiChu': khach.GhiChu,
            'ThoiGianKetThuc': khach.ThoiGianKetThuc  # Thêm ThoiGianKetThuc vào kết quả
        })
        print(khach.ThoiGianBatDau,khach.ThoiGianKetThuc)

    return result

def get_khach_dang_tiep_don_qncn():
    # Lấy danh sách KhachSiQuanModel có ThoiGianKetThuc là None
    khach_dang_tiep = KhachQNCNModel.KhachQNCNModel.objects.filter(ThoiGianKetThuc__isnull=True).order_by("-ThoiGianBatDau")
    
    # Chuẩn bị danh sách kết quả
    result = []

    for khach in khach_dang_tiep:
        # Lấy các thông tin cần thiết
        qncn = khach.QNCN
        don_vi = qncn.DonVi
        khach_hang = khach.Khach
        the_khach = khach.TheKhach

        # Thêm vào danh sách kết quả
        result.append({
            'TenSiQuan': qncn.HoTen,
            'TenDonViSiQuan': don_vi.MaDonVi,
            'TenKhach': khach_hang.HoTenKhach,
            'SoDinhDanh': khach_hang.SoDinhDanh,
            'ThoiGianBatDau': khach.ThoiGianBatDau,
            'SoTheKhach': the_khach.SoThe,
            'GhiChu': khach.GhiChu
        })

    return result


def get_khach_qncn():


# Lấy thời gian 16h30 của ngày hôm trước
    now = timezone.now()
    start_of_yesterday = (now - timedelta(days=1)).replace(hour=16, minute=30, second=0, microsecond=0)
    
    # Lọc các bản ghi có ThoiGianKetThuc là None hoặc ThoiGianBatDau từ 16h30 ngày hôm trước
    khach_dang_tiep = KhachQNCNModel.KhachQNCNModel.objects.filter(ThoiGianKetThuc__gte=start_of_yesterday).order_by("-ThoiGianBatDau")
    # khach_dang_tiep = KhachSiQuanModel.KhachSiQuanModel.objects.all().order_by("-ThoiGianBatDau")
    # khach_dang_tiep = KhachSiQuanModel.KhachSiQuanModel.objects.filter(ThoiGianKetThuc__isnull=False).order_by("-ThoiGianBatDau")
    # Chuẩn bị danh sách kết quả
    result = []

    for khach in khach_dang_tiep:
        # Lấy các thông tin cần thiết
        qncn = khach.QNCN
        don_vi = qncn.DonVi
        khach_hang = khach.Khach
        the_khach = khach.TheKhach

        # Thêm vào danh sách kết quả
        result.append({
            'TenSiQuan': qncn.HoTen,
            'TenDonViSiQuan': don_vi.MaDonVi,
            'TenKhach': khach_hang.HoTenKhach,
            'SoDinhDanh': khach_hang.SoDinhDanh,
            'ThoiGianBatDau': khach.ThoiGianBatDau,
            'ThoiGianKetThuc': khach.ThoiGianKetThuc,
            'SoTheKhach': the_khach.SoThe,
            'GhiChu': khach.GhiChu,
            'ThoiGianKetThuc': khach.ThoiGianKetThuc  # Thêm ThoiGianKetThuc vào kết quả
        })
        print(khach.ThoiGianBatDau,khach.ThoiGianKetThuc)

    return result


def get_khach_dang_tiep_don_vc():
    # Lấy danh sách KhachSiQuanModel có ThoiGianKetThuc là None
    khach_dang_tiep = KhachVienChucModel.KhachVienChucModel.objects.filter(ThoiGianKetThuc__isnull=True).order_by("-ThoiGianBatDau")
    
    # Chuẩn bị danh sách kết quả
    result = []

    for khach in khach_dang_tiep:
        # Lấy các thông tin cần thiết
        vc = khach.VienChuc
        don_vi = vc.DonVi
        khach_hang = khach.Khach
        the_khach = khach.TheKhach

        # Thêm vào danh sách kết quả
        result.append({
            'TenSiQuan': vc.HoTen,
            'TenDonViSiQuan': don_vi.MaDonVi,
            'TenKhach': khach_hang.HoTenKhach,
            'SoDinhDanh': khach_hang.SoDinhDanh,
            'ThoiGianBatDau': khach.ThoiGianBatDau,
            'SoTheKhach': the_khach.SoThe,
            'GhiChu': khach.GhiChu
        })

    return result


def get_khach_vc():


# Lấy thời gian 16h30 của ngày hôm trước
    now = timezone.now()
    start_of_yesterday = (now - timedelta(days=1)).replace(hour=16, minute=30, second=0, microsecond=0)
    
    # Lọc các bản ghi có ThoiGianKetThuc là None hoặc ThoiGianBatDau từ 16h30 ngày hôm trước
    khach_dang_tiep = KhachVienChucModel.KhachVienChucModel.objects.filter(ThoiGianKetThuc__gte=start_of_yesterday).order_by("-ThoiGianBatDau")
    # khach_dang_tiep = KhachSiQuanModel.KhachSiQuanModel.objects.all().order_by("-ThoiGianBatDau")
    # khach_dang_tiep = KhachSiQuanModel.KhachSiQuanModel.objects.filter(ThoiGianKetThuc__isnull=False).order_by("-ThoiGianBatDau")
    # Chuẩn bị danh sách kết quả
    result = []

    for khach in khach_dang_tiep:
        # Lấy các thông tin cần thiết
        vc = khach.VienChuc
        don_vi = vc.DonVi
        khach_hang = khach.Khach
        the_khach = khach.TheKhach

        # Thêm vào danh sách kết quả
        result.append({
            'TenSiQuan': vc.HoTen,
            'TenDonViSiQuan': don_vi.MaDonVi,
            'TenKhach': khach_hang.HoTenKhach,
            'SoDinhDanh': khach_hang.SoDinhDanh,
            'ThoiGianBatDau': khach.ThoiGianBatDau,
            'ThoiGianKetThuc': khach.ThoiGianKetThuc,
            'SoTheKhach': the_khach.SoThe,
            'GhiChu': khach.GhiChu,
            'ThoiGianKetThuc': khach.ThoiGianKetThuc  # Thêm ThoiGianKetThuc vào kết quả
        })
        print(khach.ThoiGianBatDau,khach.ThoiGianKetThuc)

    return result

